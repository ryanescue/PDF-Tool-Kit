from __future__ import annotations

import os
import re
import tempfile
import textwrap #text cleaner 
from dataclasses import dataclass #tracking Temp files
from io import BytesIO #buffer
from pathlib import Path
from typing import List, Sequence

from django.core.files.uploadedfile import UploadedFile #django forms file upload
from docx import Document as DocxDocument #docx reader
from PIL import Image #image formats for png and jpg
from PyPDF2 import PdfReader, PdfWriter 
from reportlab.lib.pagesizes import LETTER  #rendering plain text + Doc/docx to pdf
from reportlab.pdfgen import canvas # ^^

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png"}
TEXT_EXTENSIONS = {".txt"}
DOCX_EXTENSIONS = {".docx"}
DOC_EXTENSIONS = {".doc"}
SUPPORTED_EXTENSIONS = IMAGE_EXTENSIONS | TEXT_EXTENSIONS | DOCX_EXTENSIONS | DOC_EXTENSIONS | {".pdf"}


@dataclass
class MergeFile:#helper for file metadata
    path: Path
    original_name: str
    def cleanup(self) -> None:
        try:
            os.unlink(self.path)
        except OSError:
            pass


def merge_uploaded_files(files: Sequence[UploadedFile], *, output_label: str) -> Path:
    #merges 'files' in order 1-x
    #each 'UploadedFile' is written to a temp file
    if not files: #check
        raise ValueError("Pleae select at least one file before merging.")

    writer=PdfWriter()
    saved_files: List[MergeFile]= []
    try:
        for uploaded in files:
            temp_file= _persist_upload(uploaded)
            saved_files.append(temp_file)
            _append_to_writer(writer, temp_file)
    finally:
        #this always deletes the temp source files
        for saved in saved_files:
            saved.cleanup()
    suffix= ".pdf" if not output_label.lower().endswith(".pdf") else ""
    output_file= tempfile.NamedTemporaryFile(delete=False, suffix=suffix or ".pdf")
    try:
        writer.write(output_file)
    finally:
        output_file.close()

    return Path(output_file.name)


def _persist_upload(uploaded: UploadedFile) -> MergeFile: #write the uploaded chunk to disk to allow PyPDF/ReportLab to be able to read it
    suffix = Path(uploaded.name).suffix or ""
    temp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    for chunk in uploaded.chunks():
        temp.write(chunk)
    temp.flush()
    temp.close()
    return MergeFile(Path(temp.name), uploaded.name)

#takes file thats been saved to disk (merge_file) & adds its pages to the PdfWriter final merged PDF
def _append_to_writer(writer: PdfWriter, merge_file: MergeFile) -> None: #normalize into PDF pages and add them to the writer
    suffix = merge_file.path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file '{merge_file.original_name}'. "
            "Allowed types: PDF, DOC, DOCX, TXT, JPG, JPEG, PNG." #i wana excel files but was struggleing on size portions
        )
    #depends on suffix
    if suffix == ".pdf": 
        reader=PdfReader(str(merge_file.path)) #pdfReader
    elif suffix in IMAGE_EXTENSIONS:#.jpg,png, etc
        reader=PdfReader(_image_to_pdf_stream(merge_file.path))
    elif suffix in DOCX_EXTENSIONS:
        text=_extract_docx_text(merge_file.path)#extract
        reader=PdfReader(_text_to_pdf_stream(text, merge_file.original_name))#read extraction
    elif suffix in DOC_EXTENSIONS:
        text=_extract_doc_text(merge_file.path)#^ 
        reader=PdfReader(_text_to_pdf_stream(text, merge_file.original_name))#^
    elif suffix in TEXT_EXTENSIONS:#.txt ez
        text=merge_file.path.read_text(encoding="utf-8", errors="ignore")
        reader=PdfReader(_text_to_pdf_stream(text, merge_file.original_name))
    else:  # pragma: no cover - defensive, SUPPORTED_EXTENSIONS gate prevents this
        raise ValueError(f"File type '{suffix}' is not supported.")
    #when PdfReader itll loop through reader.pages and adds each page to the master writer
    for page in reader.pages:
        writer.add_page(page)


def _image_to_pdf_stream(path: Path) -> BytesIO: #convert an image on disk into an in-memory PDF buffer
    image = Image.open(path) #opens via Pillow
    if image.mode in ("RGBA", "P"): #converts
        image = image.convert("RGB")
    stream = BytesIO() 
    image.save(stream, format="PDF", resolution=300.0) #pdf at 300dpi
    stream.seek(0)
    return stream


def _text_to_pdf_stream(text: str, label: str) -> BytesIO: #I really struggled on this porttion as it was causing an runtime error
    cleaned= text.strip() or "[Empty document]"
    buffer= BytesIO()
    pdf= canvas.Canvas(buffer, pagesize=LETTER)
    pdf.setTitle(label)
    width, height= LETTER
    left_margin= 54
    right_margin= width - 54
    top_margin= height - 72
    y = top_margin
    line_height= 14

    def _draw_line(line: str, cursor_y: int) -> int:
        pdf.drawString(left_margin, cursor_y, line)
        return cursor_y - line_height

    wrap_width= int((right_margin - left_margin) / 7) #wraps each input line with a FIXED width and line by line
    for raw_line in cleaned.splitlines():
        wrapped = textwrap.wrap(raw_line, width=wrap_width) or [""]
        for line in wrapped:
            if y <= 72:
                pdf.showPage()
                pdf.setTitle(label)
                y = top_margin
            y = _draw_line(line, y)

    pdf.save()
    buffer.seek(0)
    return buffer


def _extract_docx_text(path: Path) -> str:#Read paragraphs + tables from a docx into a plain text
    document= DocxDocument(str(path))
    parts: List[str]= []
    for paragraph in document.paragraphs:
        parts.append(paragraph.text)
    table_text= _docx_tables_to_text(document)
    if table_text:
        parts.append(table_text)
    return "\n".join(parts)


def _docx_tables_to_text(document: DocxDocument) -> str: #flattens docx table rows into pipe lines
    rows: List[str]= []
    for table in document.tables:
        for row in table.rows:
            cells= [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cell for cell in cells if cell))
        rows.append("")
    return "\n".join(row for row in rows if row).strip()


def _extract_doc_text(path: Path) -> str: #old doc files need some love toooo, the docx extractor wasnt working so had to implement this
    #binary Word documents are tricky to parse without external utilities, i had to ask AI deeply to help me understand how to do this one
    #this method focuses on extracting readable ASCII/UTF encoded strings
    data = path.read_bytes()
    #try a couple of decoders before falling back to ascii-ish filtering
    text: str
    for encoding in ("utf-16", "cp1252", "latin-1"): #loops over common encodings used by word
        try:
            text = data.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        filtered = (chr(b) if 32 <= b <= 126 else " " for b in data)
        text = "".join(filtered)
    #text cleaning after proccesing
    text = text.replace("\r", "\n")
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()
